#!/usr/bin/env python3
from __future__ import annotations

import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WORK = ROOT / "shangrao_fault_output" / "_work"
TEMPLATE = Path("/Users/yipang/.codex/skills/company-word-docx/assets/company-template.docx")
OUTPUT = ROOT / "shangrao_fault_output" / "上饶银行故障处理报告_Codex生成.docx"
LOGO = WORK / "company_logo.png"

TITLE = "上饶银行支付系统业务中断故障处理报告"
SUBTITLE = "2016年3月2日 | 故障处理、原因分析与后续建议"
SERVICE_UNIT = "北京银信长远科技股份有限公司"
REPORT_DATE = "2026年05月03日"

BLUE = "006FC9"
DARK = "123B56"
TEXT = "1F2937"
TEAL = "19A5BE"


OUTLINE: list[tuple[int, str, int]] = [
    (1, "1. 报告摘要", 1),
    (2, "1.1 结论与当前状态", 1),
    (3, "1.1.1 报告口径说明", 2),
    (1, "2. 基本信息", 3),
    (2, "2.1 事件基础信息", 3),
    (3, "2.1.1 资料来源与证据等级", 4),
    (1, "3. 故障现象与影响", 5),
    (2, "3.1 客户侧现象", 5),
    (3, "3.1.1 受影响业务范围", 6),
    (1, "4. 处理时间线", 7),
    (2, "4.1 关键时间线", 7),
    (3, "4.1.1 恢复时间口径", 9),
    (1, "5. 排查与处理过程", 10),
    (2, "5.1 排查处置阶段", 10),
    (3, "5.1.1 关键操作与结果", 11),
    (1, "6. 原因分析", 12),
    (2, "6.1 技术原因", 12),
    (3, "6.1.1 直接原因、触发条件与贡献因素", 13),
    (2, "6.2 流程与管理因素", 14),
    (3, "6.2.1 已排除或证据不足事项", 15),
    (1, "7. 恢复验证与客户确认", 16),
    (2, "7.1 验证结果", 16),
    (3, "7.1.1 业务、资金与舆情确认", 17),
    (1, "8. 后续建议方案", 18),
    (2, "8.1 整改建议", 18),
    (3, "8.1.1 优先级与实施窗口", 19),
    (1, "9. 附录：证据索引与补充材料", 20),
    (2, "9.1 证据索引", 20),
    (3, "9.1.1 源文件清单", 21),
    (2, "9.2 口径差异说明", 22),
    (3, "9.2.1 本报告采用口径", 22),
]


OVERVIEW_ROWS = [
    ("客户名称", "上饶银行股份有限公司"),
    ("事件主题", "支付系统及相关存储卷在 SVC 磁盘镜像同步期间出现性能劣化、部分卷脱机及业务中断"),
    ("发生地点", "江西省上饶市紫阳大道佳利商城上饶银行数据中心/信息中心机房"),
    ("客户侧异常开始", "2016年3月2日12:26-12:30 左右出现 MQ 队列写入缓慢、掌易行汇款超时/卡顿等现象"),
    ("正式中断口径", "2016年3月2日13:30 至 15:20，支付系统完全中断 1小时50分"),
    ("技术恢复口径", "15:20 支付系统可访问并受理跨行汇款；16:30 多数业务系统恢复；18:10 服务商报告记录所有业务系统恢复"),
    ("影响范围", "正式报送材料确认人行支付系统业务中断，商圈和 OA 办公系统不可用；服务商材料另记录手机银行、同城清算、公务卡、支付宝等相关存储承载系统受性能影响"),
    ("当前状态", "故障已恢复；镜像同步进程已停止，主存储访问恢复，业务经验证后恢复对外服务"),
    ("结论强度", "证据链较完整。SVC 审计日志、日志分析报告、监管报送材料和处理报告可相互印证主要技术过程"),
]


TIMELINE_ROWS = [
    ("2016-02-29 16:18", "SVC 审计日志", "对测试卷执行 addvdiskcopy，进行磁盘镜像测试。", "测试完成但代表性不足，后续变更仍按日间执行。", "E08"),
    ("2016-02-29 16:29", "SVC 审计日志", "对测试卷修改镜像同步速率为 90。", "形成“白天可执行、影响不大”的判断依据之一。", "E08"),
    ("2016-03-01", "变更申请表", "提交 SVC 存储磁盘镜像变更申请，计划 2016-03-02 执行。", "申请材料称已测试评估且不会对业务造成大影响。", "E07"),
    ("2016-03-02 10:50-12:10", "SVC 审计日志", "新增 V7000_3 相关存储池、mdisk，并连续创建大量 vdisk 镜像拷贝。", "镜像同步任务进入生产存储环境。", "E08"),
    ("2016-03-02 12:22", "日志分析/监管报告", "科技人员开始/继续进行存储镜像操作。", "进入客户业务时段内的生产变更窗口。", "E03/E04"),
    ("2016-03-02 12:26", "监管报告", "人行前置机日志出现 MQ 接收队列写入缓慢、业务受理迟钝。", "支付链路开始出现性能劣化迹象。", "E04"),
    ("2016-03-02 12:30", "监管报告/处理报告", "客户通过掌易行汇款出现超时、卡顿；科技部收到反馈。", "客户可感知异常出现。", "E01/E04"),
    ("2016-03-02 12:40-12:50", "日志分析/监管报告", "检查人行前置机、支付数据库和 Oracle 日志切换，发现存储读写显著变慢。", "初步判断为镜像同步占用 SVC IO 资源导致性能下降。", "E03/E04"),
    ("2016-03-02 12:49", "SVC 审计日志", "执行 chvdisk -syncrate 50 调整部分卷同步速率。", "尝试降低同步压力，但未能消除 IO 堵塞。", "E08"),
    ("2016-03-02 13:10", "日志分析/监管报告", "尝试通过删除镜像拷贝关系终止同步进程。", "因系统资源严重不足，命令超时或未有效执行。", "E03/E04"),
    ("2016-03-02 13:30-13:34", "日志分析/监管报告/SVC 审计日志", "尝试去除镜像存储 mapping，并执行 rmvdiskcopy。", "镜像相关 vdisk 进入脱机状态，支付系统业务中断。", "E03/E04/E08"),
    ("2016-03-02 14:00-14:40", "处理报告/日志分析", "联系 IBM 实验室、IBM800 与二线专家；收集 support 信息；多次执行 recovervdisk。", "recovervdisk 未能使目标 vdisk 联机。", "E01/E03"),
    ("2016-03-02 15:00-15:08", "日志分析/SVC 审计日志", "按 IBM 专家建议重新发现/恢复存储 mapping，执行 detectmdisk，并对测试卷迁移验证。", "重新发现存储后卷自动联机，主存储数据访问恢复。", "E03/E08"),
    ("2016-03-02 15:20", "监管报告/日志分析", "主存储访问恢复，经系统验证后对外提供服务。", "各渠道受理多笔跨行汇款业务，正式支付系统中断口径结束。", "E03/E04"),
    ("2016-03-02 16:10-16:30", "服务商处理报告", "重新 mapping 后继续逐一删除 SVC 镜像卷，系统管理员重启受影响业务系统。", "大部分业务系统恢复正常对外服务。", "E01"),
    ("2016-03-02 17:00-18:10", "服务商处理报告", "发现数据整合数据库主机无法访问，两台数据库主机自动关闭；重新启动数据库主机。", "服务商报告记录所有业务系统于 18:10 恢复。", "E01"),
    ("2016-03-02 晚间至03-08 08:00", "监管报告", "进行人民银行支付系统对账，并持续观察投诉和舆情。", "头寸对平，未给客户资金带来损失；未出现客户投诉和不利舆情传播。", "E04"),
]


ACTION_ROWS = [
    ("性能确认", "检查人行前置机日志、支付数据库和 Oracle 日志切换。", "确认 MQ 队列写入缓慢、日志切换迟缓，问题指向存储 IO。", "E03/E04"),
    ("中止同步", "尝试删除镜像拷贝关系，降低或停止镜像复制。", "在资源严重不足时命令超时，标准恢复手段未能及时释放 IO。", "E03/E08"),
    ("错误处置放大", "尝试去除镜像存储 mapping。", "SVC 上镜像相关 vdisk 变为脱机，支付系统从迟缓发展为中断。", "E01/E04/E08"),
    ("厂商升级", "联系 IBM 实验室、IBM800、二线及德国三线专家，搭建远程桌面。", "形成重新 mapping 与逐一删除镜像卷的处理路径。", "E01/E04"),
    ("恢复联机", "重新将去除 mapping 的存储加入 SVC，执行 detectmdisk。", "卷自动联机，主存储数据恢复访问。", "E03/E08"),
    ("恢复业务", "按容量/风险逐一删除镜像卷，重启受影响业务系统并执行验证。", "15:20 支付系统恢复访问；16:30 多数系统恢复；18:10 服务商记录全部恢复。", "E01/E03/E04"),
]


CAUSE_ROWS = [
    ("直接原因", "SVC 磁盘镜像同步期间批量 vdisk copy 占用大量 IO 资源，导致支付链路承载存储读写性能显著下降；在同步未正常停止的情况下去除镜像存储 mapping，引发镜像相关卷脱机。"),
    ("触发条件", "2016-03-02 日间生产业务时段执行主存储向灾备存储的批量镜像同步，且同一时间段并发镜像卷数量过多。正式报送材料记载同时对 21 个 vdisk 操作，SVC auditlog 也显示 10:50-12:10 间连续创建大量 addvdiskcopy。"),
    ("贡献因素一", "应急预案对资源耗尽、命令超时、删除镜像关系失败、mapping 去除后的 vdisk 状态变化等复杂异常场景覆盖不足。"),
    ("贡献因素二", "变更前测试代表性不足。2月29日测试主要针对测试卷，未充分模拟生产数据量、并发镜像数量、业务高峰 IO 模型及支付系统 2GB 光纤通道卡约束。"),
    ("贡献因素三", "变更窗口选择不当。材料显示变更申请认为日间影响不大，但故障后整改要求明确禁止对外营业期间执行生产系统变更。"),
    ("贡献因素四", "监控体系不完整。监管报告指出资源整合过渡阶段存储监控未部署，未能在 IO 资源异常时提前预警。"),
]


RECOMMENDATION_ROWS = [
    ("P1", "立即固化生产变更冻结规则", "禁止营业时段对生产存储、支付链路和核心系统执行高 IO、高风险变更；紧急变更需单独审批并设置回退条件。", "立即执行/持续生效", "客户科技部、服务商、外包管理"),
    ("P1", "重做镜像同步实施方案和回退预案", "按卷容量、业务重要性和存储池隔离情况分批实施；明确命令超时、同步异常、mapping 误操作后的可执行恢复步骤。", "下一次变更前", "服务商主责，客户审批，厂商复核"),
    ("P1", "开展 SVC/V7000 变更前压力验证", "在测试环境或低风险窗口模拟生产级数据量、并发卷数量和支付系统 IO 模型，形成可量化阈值。", "下一维护窗口前", "服务商、IBM/原厂、客户联合"),
    ("P2", "补齐存储与支付链路监控", "建立 SVC IO 延迟、队列深度、vdisk copy 同步状态、MQ 队列、Oracle 日志切换耗时等指标监控和告警。", "1-2周", "客户运维、服务商"),
    ("P2", "建立分级应急指挥与厂商升级机制", "在重大生产故障时明确现场指挥、操作授权、原厂升级、远程接入和客户侧沟通节奏。", "1个月内", "客户科技部、服务商"),
    ("P3", "完善外包服务质量复盘机制", "将风险评估质量、预案完整性、测试代表性和变更执行纪律纳入外包考核，降低类似变更风险。", "季度复盘", "客户外包管理、服务商管理层"),
]


SOURCE_ROWS = [
    ("E01", "上饶银行存储中断故障处理报告.docx", "服务商处理报告", "12:30发现异常、13:35卷脱机、16:30大部分业务恢复、18:10全部恢复；记录完整处置过程和整改措施。"),
    ("E02", "上饶银行存储中断故障处理报告1.docx", "服务商处理报告变体", "采用“存储2控制器不稳定”表述，与日志分析和监管报告主口径不完全一致，作为差异材料引用。"),
    ("E03", "上饶银行存储中断故障日志分析报告*.docx", "日志分析报告", "记录 12:22 至 15:20 事件经过、SVC 命令清单和审计日志截图说明。"),
    ("E04", "（16）关于上饶银行支付系统业务中断事件的报告-银监局.doc", "正式监管报送材料", "采用上银文〔2016〕16号口径，确认支付系统 13:30-15:20 中断 1小时50分、影响范围、原因和整改措施。"),
    ("E05", "关于上饶银行支付系统业务中断事件的报告20160304.doc", "监管报告过程版本", "包含 2016-03-04 版本文字，部分措辞仍有占位内容，作为版本比对材料。"),
    ("E06", "关于上饶银行支付系统业务中断事件的报告-银监局*.doc", "监管报告草稿/修订件", "与正式文号材料大体一致，时间线和整改措施用于交叉核验。"),
    ("E07", "系统变更申请表_SVC磁盘镜像.docx", "变更申请", "记录 2016-03-01 申请、2016-03-02 变更日期、变更理由、影响系统和测试评估判断。"),
    ("E08", "snap.rar / SVC auditlog", "SVC 原始诊断包", "展开后核到 2月29日测试、3月2日批量 addvdiskcopy、chvdisk、rmvdiskcopy、recovervdisk、detectmdisk 等操作。"),
    ("E09", "截图解释.pptx", "日志截图说明", "说明创建磁盘镜像、修改同步速率、删除镜像配对关系、修复脱机卷、重新发现 mapping、迁移测试等截图主题。"),
]


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run._element.rPr.rFonts.set(qn("w:ascii"), "微软雅黑")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, header: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, 10.5, bold=header, color="FFFFFF" if header else TEXT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", "80"), ("bottom", "80"), ("left", "120"), ("right", "120")):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        set_cell_text(cell, header, header=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 18 and idx != len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], str(value), align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_body_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 12, bold=True, color=TEXT)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, 12, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, 12, color=TEXT)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_font(p, size={1: 18, 2: 16, 3: 15}.get(level, 14), bold=True, color=DARK if level == 1 else TEXT)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "微软雅黑")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "微软雅黑")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for level, size in ((1, 18), (2, 16), (3, 15), (4, 14)):
        style = styles[f"Heading {level}"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style._element.rPr.rFonts.set(qn("w:ascii"), "微软雅黑")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DARK if level == 1 else TEXT)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def clear_template_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_cover(doc: Document) -> None:
    with zipfile.ZipFile(TEMPLATE) as zf:
        LOGO.write_bytes(zf.read("word/media/image1.png"))

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.first_line_indent = Pt(0)
    run_logo = p_logo.add_run()
    run_logo.add_picture(str(LOGO), width=Inches(1.65))

    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    r = p_title.add_run(TITLE)
    set_run_font(r, 26, bold=True, color=DARK)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.first_line_indent = Pt(0)
    r = p_sub.add_run(SUBTITLE)
    set_run_font(r, 14, color=TEAL)

    for _ in range(7):
        doc.add_paragraph()

    meta = [
        ("服务单位", SERVICE_UNIT),
        ("服务对象", "上饶银行股份有限公司"),
        ("报告日期", REPORT_DATE),
        ("报告性质", "对外故障处理报告"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in meta:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], value, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(4.6)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("目录", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_font(p, 18, True, DARK)

    for level, text, page in OUTLINE:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt((level - 1) * 18)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))
        run = p.add_run(f"{text}\t{page}")
        set_run_font(run, 11, color=TEXT)


def build_body(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "1. 报告摘要", 1)
    add_heading(doc, "1.1 结论与当前状态", 2)
    add_body_paragraph(
        doc,
        "结合现有监管报送材料、服务商故障处理报告、日志分析报告、变更申请表以及 SVC 原始 auditlog，本次故障的主链路可以明确还原：在上饶银行 IT 基础架构资源整合项目中，生产主存储向灾备存储进行 SVC 磁盘镜像同步；批量 vdisk copy 占用了大量 SVC IO 资源，引发支付链路存储读写性能下降；随后在同步进程未能通过标准删除操作正常终止的情况下，现场尝试去除镜像存储 mapping，导致相关 vdisk 脱机，支付系统业务中断。",
    )
    add_body_paragraph(
        doc,
        "故障已经恢复。正式报送口径记录 2016年3月2日13:30 至 15:20 支付系统完全中断 1小时50分；15:20 后支付系统可访问并经验证受理跨行汇款业务。服务商处理报告同时记录了更早的性能劣化、更多存储承载业务系统影响，以及 16:30/18:10 的技术恢复里程碑。本报告将业务侧正式口径与技术侧处置口径分别列示。",
    )
    add_table(doc, ["项目", "内容"], OVERVIEW_ROWS, widths=[1.55, 5.6])

    add_heading(doc, "1.1.1 报告口径说明", 3)
    add_body_paragraph(
        doc,
        "本报告采用“正式监管报送材料 + 日志分析报告 + SVC auditlog”作为主证据链。对于服务商处理报告中“12:30-16:30影响约4小时”“18:10所有业务系统恢复”的记录，作为技术处置和扩大影响范围参考；对于另一份处理报告中“存储2控制器不稳定”的表述，由于与日志分析报告及 auditlog 主链路不一致，列为差异口径，不作为本报告根因主结论。",
    )

    add_heading(doc, "2. 基本信息", 1)
    add_heading(doc, "2.1 事件基础信息", 2)
    base_rows = [
        ("事件名称", "上饶银行支付系统业务中断及 SVC 存储镜像同步故障"),
        ("发生日期", "2016年3月2日"),
        ("项目背景", "上饶银行 IT 基础架构资源整合项目，主存储向灾备存储进行镜像同步以提高可用性和灾备能力"),
        ("涉及平台", "IBM SVC、V7000 存储、支付前置机/支付数据库及相关业务系统"),
        ("服务商/外包单位", "银信科技/北京银信长远科技股份有限公司相关现场工程师参与处置"),
        ("厂商支持", "IBM 实验室、IBM800、IBM 二线及德国三线专家提供远程支持"),
        ("报告适用范围", "用于客户侧故障处理复盘、外部沟通、整改跟踪和后续变更治理参考"),
    ]
    add_table(doc, ["字段", "内容"], base_rows, widths=[1.7, 5.45])

    add_heading(doc, "2.1.1 资料来源与证据等级", 3)
    add_body_paragraph(
        doc,
        "本报告已读取 fault_context 目录下所有文件，包括 DOC/DOCX 报告、PPTX 截图说明、变更申请表和 snap.rar 诊断包。证据优先级按原始日志、正式报送材料、日志分析报告、服务商处理报告、草稿/版本材料依次采用。",
    )
    add_table(doc, ["证据ID", "材料", "类别", "在本报告中的用途"], SOURCE_ROWS, widths=[0.65, 2.2, 1.25, 3.05])

    add_heading(doc, "3. 故障现象与影响", 1)
    add_heading(doc, "3.1 客户侧现象", 2)
    add_body_paragraph(
        doc,
        "客户侧最早可感知现象出现在 2016年3月2日12:26-12:30 左右。正式报告记录人行前置机日志从 12:26 开始存在 MQ 接收队列写入缓慢、业务受理迟钝；12:30 内部反馈客户通过掌易行汇款出现超时、卡顿。服务商报告也记录手机银行系统性能严重下降，二代支付系统随后出现同类问题。",
    )
    add_body_paragraph(
        doc,
        "在 13:30 左右去除镜像存储 mapping 后，SVC 主机上镜像相关卷进入脱机状态，加载在存储卷上的支付系统数据无法访问，支付业务由性能迟缓发展为完全中断。",
    )

    add_heading(doc, "3.1.1 受影响业务范围", 3)
    impact_rows = [
        ("正式业务影响", "人行支付系统业务中断；商圈和 OA 办公系统无法使用。", "监管报送材料 E04"),
        ("客户可感知现象", "掌易行跨行汇款超时、卡顿；客户汇款业务受影响。", "监管报送材料 E04"),
        ("未受影响/已控制事项", "柜面业务正常；急需汇款客户通过现金取款至其他银行汇款等方式进行解释和替代处理。", "监管报送材料 E04"),
        ("技术影响范围参考", "服务商报告列示手机银行、二代支付、同城清算、公务卡、支付宝、柜面、信贷等业务系统，表明相关存储承载系统存在更宽技术影响可能。", "服务商报告 E01"),
        ("资金与舆情", "3月2日晚与人民银行支付系统对账，头寸对平；截至3月8日8时未出现客户投诉和不利舆情传播。", "监管报送材料 E04"),
    ]
    add_table(doc, ["影响类别", "说明", "证据"], impact_rows, widths=[1.25, 4.45, 1.45])

    add_heading(doc, "4. 处理时间线", 1)
    add_heading(doc, "4.1 关键时间线", 2)
    add_body_paragraph(
        doc,
        "以下时间线合并正式报送材料、服务商处理报告、日志分析报告和 SVC auditlog。对客户沟通时建议优先使用正式业务影响口径，同时保留技术处置里程碑用于内部整改与服务复盘。",
    )
    add_table(doc, ["时间", "来源", "事件/动作", "结果", "证据"], TIMELINE_ROWS, widths=[1.1, 1.25, 2.55, 1.85, 0.4])

    add_heading(doc, "4.1.1 恢复时间口径", 3)
    recovery_rows = [
        ("客户业务侧正式口径", "2016年3月2日13:30-15:20", "支付系统完全中断 1小时50分；15:20 后系统可访问并受理跨行汇款。"),
        ("性能劣化起点", "2016年3月2日12:26-12:30", "MQ 队列写入缓慢、掌易行汇款超时/卡顿，属于客户可感知性能异常阶段。"),
        ("服务商技术恢复口径", "2016年3月2日16:30/18:10", "16:30 大部分业务系统恢复；18:10 服务商报告记录所有业务系统恢复。"),
    ]
    add_table(doc, ["口径", "时间", "说明"], recovery_rows, widths=[1.55, 1.55, 4.05])

    add_heading(doc, "5. 排查与处理过程", 1)
    add_heading(doc, "5.1 排查处置阶段", 2)
    add_body_paragraph(
        doc,
        "故障处置经历了性能确认、尝试中止镜像同步、异常放大、厂商升级、恢复联机、逐一清理镜像卷和业务验证六个阶段。关键转折点在于 13:30 左右去除镜像存储 mapping 后卷脱机，以及 15:01 左右重新发现存储 mapping 后卷自动联机。",
    )
    add_table(doc, ["阶段", "处置动作", "结果", "证据"], ACTION_ROWS, widths=[1.05, 2.35, 2.45, 1.3])

    add_heading(doc, "5.1.1 关键操作与结果", 3)
    operation_rows = [
        ("addvdiskcopy", "2016-03-02 11:16-12:10", "SVC auditlog 显示连续创建大量 vdisk mirror copy，形成高 IO 同步任务。"),
        ("chvdisk -syncrate 50", "2016-03-02 12:49", "尝试调整同步速率，未能根本释放资源。"),
        ("rmvdiskcopy -copy 1 31", "2016-03-02 13:34", "删除镜像拷贝关系，日志分析记录命令超时/失败，随后去除 mapping。"),
        ("recovervdisk", "2016-03-02 13:53-15:40", "多次尝试修复脱机卷，未能使关键卷联机。"),
        ("detectmdisk", "2016-03-02 15:01", "重新发现存储 mapping 后，所有卷自动联机，业务卷可访问。"),
        ("migratevdisk", "2016-03-02 15:07/15:08", "通过测试卷迁移验证存储池可用。"),
        ("rmvdiskcopy 逐一删除", "2016-03-02 16:12-17:03", "按卷逐一删除镜像拷贝关系，使 SVC 停止同步并释放资源。"),
    ]
    add_table(doc, ["操作", "时间", "说明"], operation_rows, widths=[1.5, 1.35, 4.3])

    add_heading(doc, "6. 原因分析", 1)
    add_heading(doc, "6.1 技术原因", 2)
    add_body_paragraph(
        doc,
        "本次故障技术原因可归纳为“批量镜像同步造成 IO 堵塞”和“异常处置导致 vdisk 脱机”两个连续环节。前者导致支付系统迟缓，后者导致支付系统中断。该链路由监管报送材料、日志分析报告和 SVC auditlog 共同支撑。",
    )
    add_table(doc, ["原因类型", "分析"], CAUSE_ROWS, widths=[1.55, 5.6])

    add_heading(doc, "6.1.1 直接原因、触发条件与贡献因素", 3)
    add_body_paragraph(
        doc,
        "直接原因：生产存储执行批量镜像同步时，SVC IO 资源被大量占用，支付系统相关存储读写响应显著下降；在无法正常删除镜像拷贝关系的情况下，通过去除镜像存储 mapping 试图终止同步，导致相关 vdisk 脱机。",
        bold_prefix="直接原因：",
    )
    add_body_paragraph(
        doc,
        "触发条件：变更申请材料显示本次变更安排在 2016年3月2日白天执行，且变更前判断“不会对业务系统造成大的影响”。事后材料确认，该判断低估了批量镜像对存储性能的影响。",
        bold_prefix="触发条件：",
    )
    add_body_paragraph(
        doc,
        "贡献因素：正式报告指出同一时间段批量执行多个 vdisk 镜像不符合 IBM 红皮书建议的单一 vdisk 镜像原则；支付系统 P570 小型机使用 2GB 光纤通道卡，在 IO 堵塞时抢占资源能力弱于同存储上 16GB 光纤通道卡服务器；监控未部署导致异常发现和预警滞后。",
        bold_prefix="贡献因素：",
    )

    add_heading(doc, "6.2 流程与管理因素", 2)
    process_rows = [
        ("变更评估", "测试未充分覆盖生产级数据量、并发镜像数量和业务高峰场景，导致日间实施风险被低估。"),
        ("变更窗口", "生产营业时段执行高 IO 存储变更，故障后整改明确要求营业终止后按审批计划执行。"),
        ("应急预案", "预案未充分覆盖命令超时、镜像删除失败、mapping 去除后 vdisk 脱机等异常分支。"),
        ("监控告警", "资源整合过渡阶段存储监控未部署，未能及时发现系统资源异常。"),
        ("外包管理", "监管报告要求服务商严肃处理相关人员并梳理项目外包风险，说明风险评估和现场处置质量需纳入管理闭环。"),
    ]
    add_table(doc, ["因素", "说明"], process_rows, widths=[1.35, 5.8])

    add_heading(doc, "6.2.1 已排除或证据不足事项", 3)
    add_body_paragraph(
        doc,
        "关于“存储2控制器不稳定导致数据不一致”的说法，仅在一份服务商处理报告变体中出现，未在正式监管报告、日志分析报告和 SVC auditlog 主链路中形成闭合证据。因此本报告不将其作为主根因；如后续需采用该口径，应补充原厂硬件错误日志、控制器状态记录和部件更换证明。",
    )
    add_body_paragraph(
        doc,
        "关于柜面、信贷等业务是否属于正式业务中断范围，材料存在差异。正式监管报告明确柜面业务正常，服务商处理报告列示更宽存储承载影响范围。本报告在业务影响章节中按正式报送口径陈述，并将服务商材料作为技术影响范围参考。",
    )

    add_heading(doc, "7. 恢复验证与客户确认", 1)
    add_heading(doc, "7.1 验证结果", 2)
    add_body_paragraph(
        doc,
        "恢复验证主要包括三类：一是重新 mapping/detectmdisk 后 vdisk 自动联机，主存储数据可访问；二是通过测试卷迁移确认存储池可用；三是支付系统恢复访问后，经系统验证受理多笔跨行汇款业务。",
    )
    validation_rows = [
        ("存储侧验证", "重新发现存储 mapping 后所有卷自动联机；测试卷在存储池间迁移成功。", "E03/E08"),
        ("系统侧验证", "上饶银行系统管理员对受影响业务系统进行重启并验证。", "E01/E04"),
        ("业务侧验证", "15:20 后系统可访问，各渠道受理多笔跨行汇款业务。", "E04"),
        ("资金侧验证", "3月2日晚与人民银行支付系统对账，头寸对平。", "E04"),
        ("舆情/投诉", "截至3月8日8时，未出现客户投诉和不利社会舆情传播。", "E04"),
    ]
    add_table(doc, ["验证类别", "结果", "证据"], validation_rows, widths=[1.25, 4.7, 1.2])

    add_heading(doc, "7.1.1 业务、资金与舆情确认", 3)
    add_body_paragraph(
        doc,
        "正式报送材料记录，故障期间上饶银行对需办理汇款客户进行解释，并对急需汇款客户通过现金取款至其他银行汇款等方式进行临时引导。支付系统在人行关闭前恢复，等待客户的汇款业务随后受理。",
    )
    add_body_paragraph(
        doc,
        "资金核对方面，3月2日晚与人民银行支付系统对账结果为头寸对平，未给客户资金带来损失。客户投诉和舆情方面，截至3月8日8时未出现客户投诉和不利社会舆情传播。",
    )

    add_heading(doc, "8. 后续建议方案", 1)
    add_heading(doc, "8.1 整改建议", 2)
    add_body_paragraph(
        doc,
        "后续整改应围绕“变更纪律、技术预案、监控告警、厂商升级、外包管理”五条线同步推进。建议将 SVC/V7000 生产镜像同步类操作列为高风险变更，执行前必须完成可量化压力验证、分批方案、失败回退、客户沟通和原厂复核。",
    )
    add_table(doc, ["优先级", "建议", "预期效果", "建议窗口", "责任边界"], RECOMMENDATION_ROWS, widths=[0.55, 1.55, 2.25, 1.05, 1.75])

    add_heading(doc, "8.1.1 优先级与实施窗口", 3)
    add_body_paragraph(
        doc,
        "P1 建议应在下一次生产变更前完成，重点是禁止营业时段高风险变更、重做镜像同步方案和回退预案、补齐原厂复核。P2 建议应在 1-2 周至 1 个月内落地，重点是监控告警和应急指挥机制。P3 建议纳入季度外包服务质量复盘，形成长期治理闭环。",
    )

    add_heading(doc, "9. 附录：证据索引与补充材料", 1)
    add_heading(doc, "9.1 证据索引", 2)
    add_body_paragraph(
        doc,
        "本节列示本报告引用的主要证据材料。所有抽取文本、snap 展开目录和中间文件已保存在本次输出目录的 _work 子目录，便于后续复核。",
    )
    add_table(doc, ["证据ID", "材料", "类别", "关键内容"], SOURCE_ROWS, widths=[0.65, 2.2, 1.25, 3.05])

    add_heading(doc, "9.1.1 源文件清单", 3)
    inventory_rows = [
        ("系统变更申请表_SVC磁盘镜像.docx", "变更申请", "已抽取", "记录申请人、申请日期、变更日期、变更内容和测试评估判断。"),
        ("系统变更申请表.docx", "空白/模板申请表", "已抽取", "用于确认变更申请表结构。"),
        ("上饶银行存储中断故障处理报告.docx", "服务商处理报告", "已抽取", "记录 12:30-18:10 技术处置过程。"),
        ("上饶银行存储中断故障处理报告1.docx", "服务商处理报告变体", "已抽取", "与主口径存在原因差异，列为差异材料。"),
        ("上饶银行存储中断故障日志分析报告.docx", "日志分析", "已抽取", "记录事件经过和 SVC 命令说明。"),
        ("上饶银行存储中断故障日志分析报告2.docx", "日志分析副本/版本", "已抽取", "与日志分析主稿基本一致。"),
        ("上饶银行存储中断故障日志分析报告20160311.docx", "日志分析后续版本", "已抽取", "补充 2月29日测试和 3月2日操作截图说明。"),
        ("关于上饶银行支付系统业务中断事件的报告*.doc", "监管报送材料/草稿", "已抽取", "记录正式业务影响、社会影响、原因分析和整改措施。"),
        ("（16）关于上饶银行支付系统业务中断事件的报告-银监局.doc", "正式监管报送材料", "已抽取", "采用上银文〔2016〕16号最终口径。"),
        ("截图解释.pptx", "截图说明", "已抽取", "说明日志截图对应的操作主题。"),
        ("snap.rar", "SVC/V7000 诊断包", "已展开", "展开 SVC/V7000 snap 包并核验 SVC auditlog 关键操作。"),
    ]
    add_table(doc, ["源文件", "类型", "状态", "说明"], inventory_rows, widths=[2.25, 1.2, 0.8, 2.9])

    add_heading(doc, "9.2 口径差异说明", 2)
    diff_rows = [
        ("故障时间", "正式材料：13:30-15:20 支付系统中断；服务商主报告：12:30发现异常、16:30多数系统恢复、18:10全部恢复。", "区分客户业务中断、性能劣化和技术恢复，不互相替代。"),
        ("影响范围", "正式材料：支付系统、商圈、OA；服务商报告：手机银行、二代支付、同城清算、公务卡、支付宝、柜面、信贷等。", "对外业务影响以正式材料为准，技术影响范围列示参考。"),
        ("原因表述", "主口径：批量镜像同步 IO 占用和去除 mapping 导致脱机；变体报告：存储2控制器不稳定。", "采用证据链更完整的主口径，变体作为待补充原厂硬件证据事项。"),
        ("恢复完成", "正式材料：15:20 支付系统恢复；服务商报告：16:30大部分业务、18:10全部业务恢复。", "正式对外沟通使用 15:20，技术复盘保留 16:30/18:10 里程碑。"),
    ]
    add_table(doc, ["差异项", "材料差异", "本报告处理方式"], diff_rows, widths=[1.15, 3.55, 2.45])

    add_heading(doc, "9.2.1 本报告采用口径", 3)
    add_body_paragraph(
        doc,
        "本报告最终采用如下口径：本次故障为 2016年3月2日 SVC 磁盘镜像同步生产变更期间发生的存储性能劣化与支付系统业务中断事件。客户可感知性能异常自 12:26-12:30 左右出现，支付系统正式中断时间为 13:30-15:20，主要原因是批量 vdisk 镜像同步占用 SVC IO 资源，并在异常处置中去除镜像存储 mapping 引发相关 vdisk 脱机。故障经重新 mapping、detectmdisk、逐一删除镜像卷和业务验证后恢复；资金未受损，未出现客户投诉和不利舆情。",
    )


def patch_header_text(xml: str) -> str:
    return re.sub(r"<w:t>[^<]*</w:t>", f"<w:t>{TITLE}</w:t>", xml, count=1)


def ensure_child(parent: ET.Element, tag: str) -> ET.Element:
    child = parent.find(qn(tag))
    if child is None:
        child = ET.Element(qn(tag))
        parent.append(child)
    return child


def patch_sections(document_xml: str) -> str:
    ET.register_namespace("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    ET.register_namespace("r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
    root = ET.fromstring(document_xml.encode("utf-8"))
    body = root.find(qn("w:body"))
    if body is None:
        return document_xml
    sections: list[ET.Element] = []
    for p in body.findall(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is not None:
            sect = ppr.find(qn("w:sectPr"))
            if sect is not None:
                sections.append(sect)
    final_sect = body.find(qn("w:sectPr"))
    if final_sect is not None:
        sections.append(final_sect)
    if len(sections) < 2:
        return document_xml

    def configure_common(sect: ET.Element) -> None:
        pg_sz = ensure_child(sect, "w:pgSz")
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")
        pg_mar = ensure_child(sect, "w:pgMar")
        for key, value in {
            "top": "1361",
            "right": "1304",
            "bottom": "1247",
            "left": "1417",
            "header": "720",
            "footer": "720",
            "gutter": "0",
        }.items():
            pg_mar.set(qn(f"w:{key}"), value)

    def remove_refs(sect: ET.Element) -> None:
        for child in list(sect):
            local = child.tag.rsplit("}", 1)[-1]
            if local in {"headerReference", "footerReference", "pgNumType", "titlePg"}:
                sect.remove(child)

    first = sections[0]
    remove_refs(first)
    default_header = ET.Element(qn("w:headerReference"))
    default_header.set(qn("w:type"), "default")
    default_header.set(qn("r:id"), "rId5")
    first.insert(0, default_header)
    title_pg = ET.Element(qn("w:titlePg"))
    first.insert(1, title_pg)
    configure_common(first)

    for sect in sections[1:]:
        remove_refs(sect)
        header = ET.Element(qn("w:headerReference"))
        header.set(qn("w:type"), "default")
        header.set(qn("r:id"), "rId5")
        footer = ET.Element(qn("w:footerReference"))
        footer.set(qn("w:type"), "default")
        footer.set(qn("r:id"), "rId6")
        pg_num = ET.Element(qn("w:pgNumType"))
        pg_num.set(qn("w:start"), "1")
        sect.insert(0, header)
        sect.insert(1, footer)
        sect.insert(2, pg_num)
        configure_common(sect)

    return ET.tostring(root, encoding="unicode")


def patch_docx(path: Path) -> None:
    tmp = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/header1.xml":
                data = patch_header_text(data.decode("utf-8")).encode("utf-8")
            elif item.filename == "word/document.xml":
                data = patch_sections(data.decode("utf-8")).encode("utf-8")
            elif item.filename == "docProps/core.xml":
                data = patch_core_properties(data.decode("utf-8")).encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def patch_core_properties(xml: str) -> str:
    ET.register_namespace("cp", "http://schemas.openxmlformats.org/package/2006/metadata/core-properties")
    ET.register_namespace("dc", "http://purl.org/dc/elements/1.1/")
    ET.register_namespace("dcterms", "http://purl.org/dc/terms/")
    ET.register_namespace("dcmitype", "http://purl.org/dc/dcmitype/")
    ET.register_namespace("xsi", "http://www.w3.org/2001/XMLSchema-instance")
    ns = {
        "dc": "http://purl.org/dc/elements/1.1/",
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    }
    root = ET.fromstring(xml.encode("utf-8"))
    values = {
        ("dc", "title"): TITLE,
        ("dc", "subject"): "上饶银行支付系统业务中断故障处理、原因分析与后续建议",
        ("dc", "creator"): SERVICE_UNIT,
        ("cp", "keywords"): "上饶银行;故障处理报告;SVC;存储镜像;支付系统",
        ("dc", "description"): "基于 fault_context 全部材料生成的客户侧正式故障处理报告。",
        ("cp", "lastModifiedBy"): "Codex",
    }
    for (prefix, name), value in values.items():
        tag = f"{{{ns[prefix]}}}{name}"
        node = root.find(tag)
        if node is None:
            node = ET.Element(tag)
            root.append(node)
        node.text = value
    return ET.tostring(root, encoding="unicode")


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)
    clear_template_body(doc)
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)
    build_body(doc)
    doc.save(OUTPUT)
    patch_docx(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
